from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_technical_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Technical Implementation Report: PhishDefense AI Hub', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Dissertation Technical Specification - AI-Powered Adversarial Phishing Simulation")
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "The PhishDefense AI Hub is a modular framework designed to simulate the adversarial relationship between "
        "AI-driven phishing generation (Offense) and multi-layered AI detection (Defense). The system utilizes "
        "Large Language Models (LLMs) and real-world network protocols (SMTP/POP3) to provide a high-fidelity "
        "research environment."
    )
    
    # 2. Offensive Module (Red Team)
    doc.add_heading('2. Offensive Module Technical Details', level=1)
    doc.add_paragraph(
        "The Offensive Module is responsible for generating personalized spear-phishing attacks using a context-injection "
        "engine. It leverages the Llama-3-70B model via the Groq API."
    )
    
    doc.add_heading('2.1 Contextual Phishing Generation', level=2)
    doc.add_paragraph(
        "Using LangChain Expression Language (LCEL), the system dynamiclly constructs emails based on the "
        "Enron Email Dataset. It injects specific personas, professional contexts, and psychological tones "
        "(Low Urgency, High Urgency, Passive-Aggressive) into the prompt template."
    )
    
    doc.add_heading('2.2 URL Obfuscation Engine', level=2)
    doc.add_paragraph(
        "To bypass simple blacklisting, the module implements a custom obfuscator that generates three types of malicious links:"
    )
    doc.add_paragraph("• Typosquatting: Automatic character manipulation (e.g., duplicate letters).", style='List Bullet')
    doc.add_paragraph("• Subdomain Masking: Prefixing brand names with security keywords.", style='List Bullet')
    doc.add_paragraph("• Suffix Deception: Appending misleading corporate suffixes (e.g., -auth.net).", style='List Bullet')
    
    # 3. Defensive Module (Blue Team)
    doc.add_heading('3. Defensive Module Technical Details', level=1)
    doc.add_paragraph(
        "The Defensive Module employs a 'Defense-in-Depth' approach, combining legacy heuristics with modern behavioral "
        "intelligence and semantic LLM reasoning."
    )
    
    doc.add_heading('3.1 Multi-Layered Detection Engine', level=2)
    doc.add_paragraph(
        "The detection pipeline runs every email through three distinct analysis layers:"
    )
    doc.add_paragraph("1. Heuristic Layer: Scans for technical indicators of spoofing and malicious syntax.", style='List Number')
    doc.add_paragraph("2. Behavioral Layer: Analyzes the communication relationship between sender and recipient against a historical baseline.", style='List Number')
    doc.add_paragraph("3. LLM Semantic Layer: Uses AI to detect psychological pressure and social engineering intent.", style='List Number')
    
    doc.add_heading('3.2 Hybrid Risk Scoring', level=2)
    doc.add_paragraph(
        "The final 'Alert' status is determined by a weighted voting system: "
        "Combined Score = (Heuristics * 0.3) + (Behavioral * 0.3) + (LLM Semantic * 0.4)."
    )
    
    # 4. Network Protocol Bridge
    doc.add_heading('4. Network & Gateway Integration', level=1)
    doc.add_paragraph(
        "A critical component of the research is the use of real-world protocols to move beyond local mocks. "
        "The MailtrapBridge class implements standard mail handling:"
    )
    doc.add_paragraph("• Outbound: Transmission via SMTP (Simple Mail Transfer Protocol) to an internet-based gateway.", style='List Bullet')
    doc.add_paragraph("• Inbound: Retrieval via POP3 (Post Office Protocol v3) for defensive analysis.", style='List Bullet')
    
    # 5. Evaluation Methodology
    doc.add_heading('5. Scientific Evaluation & Metrics', level=1)
    doc.add_paragraph(
        "The success of the defense is measured using three core metrics:"
    )
    doc.add_paragraph("• Detection Rate: Accuracy of the ‘Alert’ status against known offensive intent.", style='List Bullet')
    doc.add_paragraph("• Analysis Method Bias: Measuring the performance delta between heuristic-only vs AI-integrated detection.", style='List Bullet')
    doc.add_paragraph("• Tone Sensitivity: Determining which emotional trigger is most effective at evading AI detection.", style='List Bullet')
    
    # Footer
    doc.add_paragraph("\n[Report Generated for Dissertation Documentation]")
    
    # Save
    report_path = "PhishDefense_Technical_Implementation.docx"
    doc.save(report_path)
    print(f"Report created: {report_path}")

if __name__ == "__main__":
    create_technical_report()
