from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_dissertation_report():
    doc = Document()

    # --- STYLE SETUP ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- TITLE ---
    title = doc.add_heading('Comprehensive Technical Implementation: Offensive Phishing Generation Module', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- OVERVIEW ---
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        "This project is structured as an automated, ethically-sandboxed pipeline for studying and generating "
        "advanced phishing attacks. By grounding LLM generation in real-world corporate data (Enron), we move "
        "beyond generic spam toward sophisticated, context-aware 'Spear Phishing'."
    )

    # --- PHASE 1 ---
    doc.add_heading('Phase 1: Data Ingestion & Grounding (The "Baseline")', level=1)
    doc.add_paragraph(
        "To create realistic phishing, the system first needs to understand what benign (normal) corporate "
        "communication looks like."
    )
    doc.add_paragraph("The Source:", style='Heading 2')
    doc.add_paragraph(
        "We use the Enron Email Dataset, which contains real internal communications from a major corporation. "
        "This provides a footprint of authentic organizational behavior."
    )
    doc.add_paragraph("The Logic (email_parser.py):", style='Heading 2')
    doc.add_paragraph("• Loads the massive 1.3GB CSV in chunks to maintain memory efficiency (RAM safety).", style='List Bullet')
    doc.add_paragraph("• Uses Python’s specialized email library to parse complex MIME-encoded message structures.", style='List Bullet')
    doc.add_paragraph("• Extracts critical 'Targeting Data': 'From' (identity), 'To' (recipient), 'Subject' (topic), and 'Body' (context).", style='List Bullet')
    doc.add_paragraph("Output:", style='Heading 2')
    doc.add_paragraph(
        "A cleaned file (processed_enron.csv) that acts as a directory of 'Target Personas' and their "
        "unique communication styles."
    )

    # --- PHASE 2 ---
    doc.add_heading('Phase 2: Prompt Engineering (The "Brain")', level=1)
    doc.add_paragraph(
        "We move beyond simple prompts. We use LangChain Prompt Templates to enforce a rigorous psychological framework."
    )
    doc.add_paragraph("The Logic (prompt_templates.py):", style='Heading 2')
    doc.add_paragraph(
        "• Spear Phishing: Injects target names, specific roles, and real topics from Enron history to create 'contextual urgency'.", 
        style='List Bullet'
    )
    doc.add_paragraph("• Brand Impersonation: Standardizes tone to mimic entities like Microsoft, Banks, or common SaaS providers.", style='List Bullet')
    doc.add_paragraph("• Internal Notifications: Leverages corporate jargon and compliance triggers (HR policy, IT updates).", style='List Bullet')
    doc.add_paragraph("Context Injection:", style='Heading 2')
    doc.add_paragraph(
        "LangChain allows for programmatic swapping of {target_name} or {context} variables at scale, ensuring "
        "no two generated emails are identical."
    )

    # --- PHASE 3 ---
    doc.add_heading('Phase 3: Generation & Model Selection (The "Execution")', level=1)
    doc.add_paragraph(
        "The system utilizes high-speed inference engines to simulate the scale of modern automated threats."
    )
    doc.add_paragraph("The Logic (phishing_generator.py):", style='Heading 2')
    doc.add_paragraph("• Initializes connection to the Groq API for ultra-low latency generation.", style='List Bullet')
    doc.add_paragraph("• Model: Utilizes Llama-3.3-70b-versatile for its superior reasoning in social engineering tasks.", style='List Bullet')
    doc.add_paragraph("• Parameterization: Sets temperature=0.7 to balance creative deception with professional structure.", style='List Bullet')
    doc.add_paragraph("Scalability:", style='Heading 2')
    doc.add_paragraph("Designed to iterate through the entire Enron database, creating thousands of unique samples automatically.")

    # --- PHASE 4 ---
    doc.add_heading('Phase 4: Automated Labelling & Logging (The "Traceability")', level=1)
    doc.add_paragraph("For research, data without metadata is inadmissible. Every email is 'fingerprinted'.")
    doc.add_paragraph("The Logic:", style='Heading 2')
    doc.add_paragraph(
        "Every generation is appended to a structured JSONL log (data/generated_phishing_v1.jsonl) containing:", 
        style='List Bullet'
    )
    doc.add_paragraph("• Model Metadata: Which exact LLM produced the content?", style='List Bullet 2')
    doc.add_paragraph("• Attack Type: Classification of the social engineering strategy used.", style='List Bullet 2')
    doc.add_paragraph("• Input Grounding: Which original Enron email served as the source?", style='List Bullet 2')
    doc.add_paragraph("• Timestamps: Accurate records of the generation session.", style='List Bullet 2')

    # --- PHASE 5 ---
    doc.add_heading('Phase 5: Ethical Sandboxing (The "Safety Valve")', level=1)
    doc.add_paragraph("This phase ensures 100% compliance with ethical research standards.")
    doc.add_paragraph("The Logic (mock_smtp.py):", style='Heading 2')
    doc.add_paragraph("• Outbound Interception: Standard SMTP calls are intercepted by a built-in Mock Handler.", style='List Bullet')
    doc.add_paragraph("• Local Delivery: Emails are 'delivered' to a local directory (data/mock_inbox/) as JSON files.", style='List Bullet')
    doc.add_paragraph("Result:", style='Heading 2')
    doc.add_paragraph(
        "Researchers can safely 'send' thousands of emails without a single packet ever leaving the local machine."
    )

    # --- ORCHESTRATION ---
    doc.add_heading('The Orchestration Flow', level=1)
    doc.add_paragraph(
        "The src/main_orchestrator.py script synchronizes these phases into a single loop:"
    )
    doc.add_paragraph("1. Select random persona from Enron.", style='List Number')
    doc.add_paragraph("2. Analyze historical context (e.g., 'Western Market Forecast').", style='List Number')
    doc.add_paragraph("3. Instruct Groq (Llama-3.3) to craft a targeted attack based on that context.", style='List Number')
    doc.add_paragraph("4. Label and log all generation metadata.", style='List Number')
    doc.add_paragraph("5. Deliver to the ethical sandbox for review.", style='List Number')

    # Save
    output_path = "Detailed_Implementation_Workflow.docx"
    doc.save(output_path)
    print(f"Comprehensive report successfully saved to {os.path.abspath(output_path)}")

if __name__ == "__main__":
    create_dissertation_report()
